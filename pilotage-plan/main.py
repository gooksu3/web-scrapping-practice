from bs4 import BeautifulSoup
from selenium import webdriver
from docx import Document
from tkinter import Tk,Frame,Label,Button,font
from tkinter.ttk import Combobox,Style
from datetime import datetime
import os
import tkinter.messagebox as msgbox


date=datetime.now()
year=date.year
month=date.month
day=date.day
hour=date.hour
weekdays=["월","화","수","목","금","토","일"]
weektoday=weekdays[datetime.today().weekday()]

def modify_word_file(results):
    path_of_word_file="C:/Users/gooks/OneDrive/문서/GitHub/web-scrapping-practice/pilotage-plan/pilotage_plan.docx"
    doc=Document(path_of_word_file)
    table=doc.tables[1]
    doc.paragraphs[2].text=f"{year} 년 {month}월 {day}일 {weektoday}요일"
    num=0
    for row in table.rows:
        if num<2:
            num+=1
        else:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text=""
    num=0
    for result in results[:30]:
        time=result["time"]
        ship_name=result["ship_name"]
        pilot_in_charge=result["pilot_in_charge"]
        callsign=result["callsign"]
        from_area=result["from_area"]
        to_area=result["to_area"]
        row=num+2
        table.rows[row].cells[0].paragraphs[0].text=ship_name
        table.rows[row].cells[1].paragraphs[0].text=callsign
        table.rows[row].cells[5].paragraphs[0].text=pilot_in_charge
        table.rows[row].cells[7].paragraphs[0].text=from_area
        table.rows[row].cells[8].paragraphs[0].text="→"
        table.rows[row].cells[9].paragraphs[0].text=to_area
        num+=1  
    doc.save("C:/Users/gooks/OneDrive/문서/GitHub/web-scrapping-practice/pilotage-plan/pilotage_plan.docx")
    return path_of_word_file

def get_pilot_plans_on_word(from_time,to_time):
    if from_time=="":
        msgbox.showinfo("알림", "from 시간을 입력해주십시오.")
    elif to_time=="":
        msgbox.showinfo("알림", "to 시간을 입력해주십시오.")
    elif int(from_time[:2])==int(to_time[:2]):
        msgbox.showinfo("알림", "from시간과 to시간이 같습니다. 시간을 확인하십시오.")
    elif int(from_time[:2])>int(to_time[:2]):
        msgbox.showinfo("알림", "from시간이 to시간보다 큽니다. 시간을 확인하십시오.")
    else:
        browser=webdriver.Chrome()
        base_url="http://www.ulsanpilot.co.kr/main/pilot_forecast.php"
        response=browser.get(base_url)
        soup=BeautifulSoup(browser.page_source,"html.parser")
        table=soup.select_one("#tug_forecast")
        today_plan=table.select(".lite-inner .move_table")[0]
        columns=today_plan.select("tr")
        results=[]
        for column in columns:
            tds=column.select("td")
            num=0
            dict_plans={}
            for td in tds:
                if num==3:
                    pilot_time=td.string
                    if int(from_time[:2])>int(pilot_time[:2]) or int(to_time[:2])<int(pilot_time[:2]):
                        break
                    else:
                        dict_plans["time"]=td.string
                elif num==4:
                    dict_plans["ship_name"]=td.string
                elif num==5:
                    pilot=td.string
                    pilot_in_charge=""
                    if pilot!="":
                        pilot_number=""
                        bool_get_number=False
                        for i in pilot:
                            if i=="(":
                                bool_get_number=True
                                continue
                            elif i==")":
                                bool_get_number=False
                            if bool_get_number:
                                pilot_number+=i
                        if len(pilot_number)>2:
                            while pilot_number!="":
                                pilot_in_charge+=pilot_number[:2]
                                pilot_number=pilot_number[2:]
                                if len(pilot_number)!=0:
                                    pilot_in_charge+=","
                        else:
                            pilot_in_charge=pilot_number                  
                    dict_plans["pilot_in_charge"]=pilot_in_charge
                elif num==6:
                    dict_plans["callsign"]=td.string
                elif num==10:
                    from_area=td.string
                    if from_area=="P/S":
                        from_area=""
                    dict_plans["from_area"]=from_area
                elif num==11:
                    to_area=td.string
                    if to_area=="P/S":
                        to_area=""
                    dict_plans["to_area"]=to_area
                num+=1
            if len(dict_plans)!=0:
                results.append(dict_plans)
        os.system(modify_word_file(results))

def label_font(size,weight):
    return font.Font(size=size,weight=weight)

root=Tk()

style=Style(root)
style.theme_use("vista")

root.title("도선예보 출력")
root.configure(background="#E3E3E3")
list_times=[]
for i in range(25):
    if i>hour:
        hour_in_list_time=str(i)
        if len(hour_in_list_time)<2:
            hour_in_list_time="0"+hour_in_list_time
        list_times.append(f"{hour_in_list_time}:00")
frame_name=Frame(root,background="#E3E3E3")
frame_name.grid(row=0,column=0)
label_name=Label(frame_name,text="도선예보 출력(워드파일)",font=label_font(20,"bold"),background="#E3E3E3")
label_name.grid(row=0,column=0,pady=20)
label_explanation=Label(frame_name,text="1. 시작시간과 종료시간을 설정\n2. 출력 버튼을 클릭!!\n3. 워드파일이 나오면 인쇄\n4. 인쇄 후 워드파일 반드시 종료\n5. 프로그램 종료\n(※재인쇄 시 재시작 필요)",font=label_font(12,"normal"),justify="left",background="#FCFCFC",highlightbackground="#010101",highlightthickness=1)
label_explanation.grid(row=1,column=0,pady=10)
frame_to_set_time=Frame(root,background="white",highlightbackground="#F90001",highlightthickness=1)
frame_to_set_time.grid(row=1,column=0)
frame_date=Frame(frame_to_set_time,background="white")
frame_date.grid(row=0,column=0)
label_date=Label(frame_date,text=f"{year}년{month}월{day}일",font=label_font(15,"bold"),background="white",fg="#010101")
label_date.grid(row=0,column=0)
frame_time=Frame(frame_to_set_time,background="white")
frame_time.grid(row=1,column=0,padx=10,pady=5)
label_from=Label(frame_time,text="From",font=label_font(13,"bold"),fg="#010101",background="white")
label_from.grid(row=0,column=0)
combo_from_time=Combobox(frame_time,width=6,value=list_times,font=label_font(13,"bold"))
combo_from_time.grid(row=0,column=1)
if "06:00" in list_times:
    combo_from_time.current("06:00")
label_to=Label(frame_time,text="to",font=label_font(13,"bold"),fg="#010101",background="white")
label_to.grid(row=0,column=2)
combo_to_time=Combobox(frame_time,width=6,value=list_times,font=label_font(13,"bold"))
combo_to_time.grid(row=0,column=3)
if combo_from_time.get()!="":
    combo_to_time.current("09:00")
frame_btns=Frame(root,background="#E3E3E3")
frame_btns.grid(row=2,column=0,pady=10)
btn_to_create_word=Button(frame_btns, text="출  력",width=8,font=label_font(12,"bold"),background="#EBEBEB")
btn_to_create_word.grid(row=0,column=0,padx=10,pady=3)
btn_to_end=Button(frame_btns, text="종  료",command=root.quit,width=8,font=label_font(12,"bold"),background="#EBEBEB")
btn_to_end.grid(row=0,column=1,padx=10,pady=3)
btn_to_create_word.configure(command=lambda:get_pilot_plans_on_word(combo_from_time.get(),combo_to_time.get()))

root.mainloop()